#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Audit formatting of the generated docx against the checklist."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('artykul/artykul_iot_security_radar.docx')

issues = []
ok = []

# ── margins ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
from docx.oxml.ns import qn

def check(cond, label, detail=""):
    if cond:
        ok.append(f"OK  {label}")
    else:
        issues.append(f"ERR {label} {detail}")

check(abs(section.left_margin.cm - 2.5) < 0.01,   "left_margin=2.5cm",   f"got {section.left_margin.cm:.2f}")
check(abs(section.right_margin.cm - 2.5) < 0.01,  "right_margin=2.5cm",  f"got {section.right_margin.cm:.2f}")
check(abs(section.top_margin.cm - 2.5) < 0.01,    "top_margin=2.5cm",    f"got {section.top_margin.cm:.2f}")
check(abs(section.bottom_margin.cm - 2.5) < 0.01, "bottom_margin=2.5cm", f"got {section.bottom_margin.cm:.2f}")
check(abs(section.gutter.cm - 0.5) < 0.01,        "gutter=0.5cm",        f"got {section.gutter.cm:.2f}")

sectPr = section._sectPr
pgMar = sectPr.find(qn("w:pgMar"))
mirror = pgMar.get(qn("w:mirrorMargins")) if pgMar is not None else None
check(mirror == "1", "mirror_margins=1", f"got {mirror}")

# ── paragraphs ───────────────────────────────────────────────────────────────
ALLOWED_FONTS = {"Times New Roman", "Courier New", None}
ALLOWED_SIZES = {9.0, 10.0, 11.0, 12.0, 13.0, 14.0, None}

bad_font_paras = []
bad_size_paras = []
body_no_indent = []

# Expected headings (should be centered, bold, 13pt)
heading_markers = [
    "WSTEP", "WSTĘP", "ZAGROŻENIA", "ARCHITEKTURA", "ZBIERANIE", "KLASYFIKACJA",
    "WERYFIKACJA", "WNIOSKI", "BIBLIOGRAFIA", "ABSTRACT",
    "1. ", "2. ", "3. ", "4. ", "5. ",
    "1.1 ", "1.2 ", "2.1 ", "2.2 ", "3.1 ", "3.2 ", "4.1 ", "4.2 ", "4.3 ", "5.1 ", "5.2 "
]

heading_paras = []
body_paras = []

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue

    pf = p.paragraph_format
    fi = pf.first_line_indent
    fi_cm = round(fi.cm, 2) if fi else 0.0
    al = pf.alignment
    sb = pf.space_before.pt if pf.space_before else 0
    sa = pf.space_after.pt if pf.space_after else 0

    is_heading_text = any(txt.startswith(m) for m in heading_markers)
    has_large_font = any(
        (run.font.size and run.font.size.pt >= 13) for run in p.runs
    )
    is_heading = is_heading_text or has_large_font

    for run in p.runs:
        fname = run.font.name
        fsize = run.font.size.pt if run.font.size else None
        if fname not in ALLOWED_FONTS:
            bad_font_paras.append(f"  Para {i}: font={fname!r} | {txt[:50]!r}")
        if fsize not in ALLOWED_SIZES:
            bad_size_paras.append(f"  Para {i}: size={fsize} | {txt[:50]!r}")

    # Body text (not heading, not list, not code) should have indent 1.27cm
    is_list = txt.startswith('\u2013') or txt.startswith('-')  # en-dash list
    is_code = 'Courier' in ''.join((run.font.name or '') for run in p.runs)
    is_num_list = len(txt) > 2 and txt[0].isdigit() and txt[1] in '.)'

    if (not is_heading and not is_list and not is_code and not is_num_list
            and al == WD_ALIGN_PARAGRAPH.JUSTIFY and len(txt) > 30):
        if abs(fi_cm - 1.27) > 0.05:
            body_no_indent.append(f"  Para {i}: fi={fi_cm:.2f}cm | {txt[:50]!r}")

    # Heading: should be centered — but only if font >= 13pt (not bibliography 12pt)
    if has_large_font and al != WD_ALIGN_PARAGRAPH.CENTER:
        heading_paras.append(f"  Para {i}: heading NOT centered | {txt[:50]!r}")

check(not bad_font_paras, "All fonts = TNR or Courier New")
if bad_font_paras:
    for x in bad_font_paras: print(x)

check(not bad_size_paras, "All sizes in {9,10,11,12,13,14}")
if bad_size_paras:
    for x in bad_size_paras: print(x)

check(not body_no_indent, "Body paragraphs have 1.27cm indent")
if body_no_indent:
    for x in body_no_indent: print(x)

check(not heading_paras, "Headings centered")
if heading_paras:
    for x in heading_paras: print(x)

# ── specific heading spacing ─────────────────────────────────────────────────
heading_spacing = {}
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in ("WSTĘP", "WNIOSKI", "BIBLIOGRAFIA", "ABSTRACT", "STRESZCZENIE"):
        sb = p.paragraph_format.space_before.pt if p.paragraph_format.space_before else 0
        sa = p.paragraph_format.space_after.pt if p.paragraph_format.space_after else 0
        heading_spacing[txt] = (sb, sa)

for name, (sb, sa) in heading_spacing.items():
    # STRESZCZENIE checklist specifies only "po 12pt", no before requirement
    # WNIOSKI: before=24pt (bigger!), after=12pt
    # ABSTRACT: before=24pt, after=12pt
    want_sa = 12
    check(abs(sa - want_sa) < 1, f"{name}: space_after=12pt", f"got {sa}pt")
    if name == "WNIOSKI":
        check(abs(sb - 24) < 1, f"WNIOSKI: space_before=24pt", f"got {sb}pt")
    elif name in ("ABSTRACT",):
        check(abs(sb - 24) < 1, f"{name}: space_before=24pt", f"got {sb}pt")
    elif name in ("WSTEP", "WSTĘP", "BIBLIOGRAFIA"):
        check(abs(sb - 12) < 1, f"{name}: space_before=12pt", f"got {sb}pt")

# ── author block ─────────────────────────────────────────────────────────────
p0 = doc.paragraphs[0]
r0 = p0.runs[0]
check(r0.font.size.pt == 11, "Author name: 11pt", f"got {r0.font.size.pt}")
check(r0.font.bold == True, "Author name: bold", f"got {r0.font.bold}")
check(p0.paragraph_format.space_before.pt == 12, "Author name: space_before=12pt")

p1 = doc.paragraphs[1]
r1 = p1.runs[0]
check(r1.font.size.pt == 11, "Affiliation: 11pt", f"got {r1.font.size.pt}")
check(r1.font.bold != True, "Affiliation: NOT bold", f"got bold={r1.font.bold}")

# ── title ────────────────────────────────────────────────────────────────────
p_title = doc.paragraphs[2]
r_title = p_title.runs[0]
check(r_title.font.size.pt == 14, "Title: 14pt", f"got {r_title.font.size.pt}")
check(r_title.font.bold == True, "Title: bold")
check(p_title.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER, "Title: centered")
check(p_title.paragraph_format.space_before.pt == 30, "Title: space_before=30pt")
check(p_title.paragraph_format.space_after.pt == 30, "Title: space_after=30pt")

# ── tables ───────────────────────────────────────────────────────────────────
check(len(doc.tables) == 2, "2 tables present", f"got {len(doc.tables)}")
for ti, tbl in enumerate(doc.tables):
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    sz = run.font.size.pt if run.font.size else None
                    if sz is not None and sz != 10.0:
                        issues.append(f"ERR Table {ti}: cell text size={sz} (want 10)")
                    elif sz == 10.0:
                        pass  # OK

# ── keywords ─────────────────────────────────────────────────────────────────
kw_label = doc.paragraphs[5]
check("Słowa kluczowe:" in kw_label.text, "Keywords label found")
r_kw = kw_label.runs[0]
check(r_kw.font.underline == True, "Keywords label: underlined", f"got {r_kw.font.underline}")

# ── summary ──────────────────────────────────────────────────────────────────
print()
print("=" * 55)
print(f"PASS: {len(ok)}")
print(f"FAIL: {len([x for x in issues if x.startswith('ERR')])}")
print()
print("FAILURES:")
for iss in issues:
    if iss.startswith("ERR"):
        print(f"  {iss}")
print()
print("PASSED:")
for o in ok:
    print(f"  {o}")

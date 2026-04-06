#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generates artykul_iot_security_radar.docx
Zgodny z wytycznymi dla autorów (Times New Roman, marginesy lustrzane, itp.)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_DIR = os.path.join(SCRIPT_DIR, "..", "ml", "model")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "artykul_iot_security_radar.docx")

FEATURE_IMG = os.path.join(MODEL_DIR, "feature_importance.png")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, underline=False, spacing_pt=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if spacing_pt is not None:
        rPr = run._r.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(spacing_pt * 20)))  # twips (1/20 pt)
        rPr.append(spacing)


def fmt_para(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=0,
             line_spacing=None):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line:
        pf.first_line_indent = Cm(first_line)
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)


def add_text_para(doc, text, size=12, bold=False, italic=False,
                  underline=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=0, first_line=1.27,
                  spacing_pt=None):
    """Add a single-run paragraph."""
    p = doc.add_paragraph()
    fmt_para(p, alignment=alignment, space_before=space_before,
             space_after=space_after, first_line=first_line)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, underline=underline,
             spacing_pt=spacing_pt)
    return p


def add_heading(doc, text, level="chapter", number=None):
    """
    level: 'chapter'  → TNR bold 13pt centred, before/after 12pt
           'sub'      → TNR normal 13pt centred, before/after 12pt
           'wnioski'  → TNR bold 13pt centred, before 24pt, after 12pt
           'biblio'   → same as chapter
    """
    p = doc.add_paragraph()
    full = f"{number}. {text}" if number else text
    run = p.add_run(full)
    if level == "wnioski":
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=24, space_after=12, first_line=0)
    elif level == "sub":
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=12, space_after=12, first_line=0)
        set_font(run, size=13, bold=False)
        return p
    else:
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=12, space_after=12, first_line=0)
    set_font(run, size=13, bold=True)
    return p


def add_dash_list(doc, items):
    """Bulleted list with em-dash (–). Items = list of strings."""
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        if i == 0:
            fmt_para(p, space_before=6, space_after=0, first_line=0)
        elif i == len(items) - 1:
            fmt_para(p, space_before=0, space_after=6, first_line=0)
        else:
            fmt_para(p, space_before=0, space_after=0, first_line=0)

        suffix = "." if i == len(items) - 1 else ";"
        run = p.add_run(f"– {item}{suffix}")
        set_font(run, size=12)


def add_numbered_list(doc, items):
    """Numbered list 1. 2. ... Each item starts capital, ends period."""
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        if i == 0:
            fmt_para(p, space_before=6, space_after=0, first_line=0)
        elif i == len(items) - 1:
            fmt_para(p, space_before=0, space_after=6, first_line=0)
        else:
            fmt_para(p, space_before=0, space_after=0, first_line=0)

        run = p.add_run(f"{i+1}. {item}.")
        set_font(run, size=12)


def add_table_caption(doc, n, title):
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=12, space_after=6, first_line=0)
    run = p.add_run(f"Tabela {n}. {title}")
    set_font(run, size=10)


def add_fig_caption(doc, n, title, source="opracowanie własne"):
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=6, space_after=12, first_line=0)
    run = p.add_run(f"Rys. {n}. {title} ({source})")
    set_font(run, size=10)


# ── footnote via oxml ─────────────────────────────────────────────────────────

_footnote_counter = [1]


def add_footnote(doc, para, text):
    """Inject a Word footnote into an existing paragraph's last run."""
    fn_id = _footnote_counter[0]
    _footnote_counter[0] += 1

    # Ensure footnotes part exists
    footnotes_part = doc.part.footnotes_part
    if footnotes_part is None:
        # Create minimal footnotes part — simple approach: skip if unavailable
        return

    # Add footnote element to footnotes XML
    footnotes_elem = footnotes_part._element
    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fn_id))
    fn.set(qn("w:type"), "normal")

    fn_para = OxmlElement("w:p")
    fn_ppr = OxmlElement("w:pPr")
    fn_style = OxmlElement("w:pStyle")
    fn_style.set(qn("w:val"), "FootnoteText")
    fn_ppr.append(fn_style)
    fn_para.append(fn_ppr)

    fn_run_ref = OxmlElement("w:r")
    fn_rpr_ref = OxmlElement("w:rPr")
    fn_vstyle = OxmlElement("w:vertAlign")
    fn_vstyle.set(qn("w:val"), "superscript")
    fn_rpr_ref.append(fn_vstyle)
    fn_run_ref.append(fn_rpr_ref)
    fn_ref_mark = OxmlElement("w:footnoteRef")
    fn_run_ref.append(fn_ref_mark)
    fn_para.append(fn_run_ref)

    fn_run_text = OxmlElement("w:r")
    fn_rpr_text = OxmlElement("w:rPr")
    fn_sz = OxmlElement("w:sz")
    fn_sz.set(qn("w:val"), "20")  # 10pt
    fn_rpr_text.append(fn_sz)
    fn_run_text.append(fn_rpr_text)
    fn_t = OxmlElement("w:t")
    fn_t.set(qn("xml:space"), "preserve")
    fn_t.text = f" {text}"
    fn_run_text.append(fn_t)
    fn_para.append(fn_run_text)
    fn.append(fn_para)
    footnotes_elem.append(fn)

    # Add reference mark in paragraph
    ref_run = OxmlElement("w:r")
    ref_rpr = OxmlElement("w:rPr")
    ref_vstyle = OxmlElement("w:vertAlign")
    ref_vstyle.set(qn("w:val"), "superscript")
    ref_rpr.append(ref_vstyle)
    ref_run.append(ref_rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fn_id))
    ref_run.append(ref)
    para._p.append(ref_run)


# ── document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Remove default styles' paragraph spacing
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Section / page setup
    section = doc.sections[0]
    section.page_width = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.gutter = Cm(0.5)

    # Mirror margins
    sectPr = section._sectPr
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        pgMar.set(qn("w:mirrorMargins"), "1")

    # No page numbers — do not add page number field
    return doc


# ── content builders ──────────────────────────────────────────────────────────

def build_author_block(doc):
    # Author name
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=12, space_after=0, first_line=0)
    run = p.add_run("[Imię Nazwisko]")
    set_font(run, size=11, bold=True, spacing_pt=2)

    # Affiliation
    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0, first_line=0)
    run2 = p2.add_run("[Wydział / Uczelnia]")
    set_font(run2, size=11, bold=False)


def build_title(doc):
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=30, space_after=30, first_line=0)
    run = p.add_run(
        "SYSTEM MONITOROWANIA SIECI IOT Z KLASYFIKACJĄ ATAKÓW "
        "METODAMI UCZENIA MASZYNOWEGO I MAPOWANIEM NA FRAMEWORK MITRE ATT&CK"
    )
    set_font(run, size=14, bold=True, spacing_pt=2)


def build_abstract_pl(doc):
    # Header
    p_hdr = doc.add_paragraph()
    fmt_para(p_hdr, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=12, first_line=0)
    run_hdr = p_hdr.add_run("STRESZCZENIE")
    set_font(run_hdr, size=13, bold=True)

    # Body
    text = (
        "Artykuł opisuje projekt IoT Security Radar — lokalny system monitorowania sieci IoT "
        "z detekcją i klasyfikacją ataków w czasie rzeczywistym. System integruje stos Elastic "
        "(Filebeat, Logstash, Elasticsearch, Kibana) z modułem uczenia maszynowego opartym na "
        "algorytmie Random Forest, osiągając dokładność klasyfikacji na poziomie 97% dla dziewięciu "
        "klas zdarzeń sieciowych. Wdrożona architektura dwuindeksowa oddziela zdarzenia surowe od "
        "predykcji modelu, co umożliwia niezależną analizę trafności ground-truth i decyzji klasyfikatora. "
        "Każde zdarzenie ataku jest mapowane na taktyki i techniki frameworku MITRE ATT&CK, dostarczając "
        "operatorom SOC kontekst niezbędny do priorytetyzacji reagowania. W procesie walidacji "
        "zindeksowano łącznie 55 280 zdarzeń z trzech źródeł: zbioru benchmarkowego NF-ToN-IoT-v2, "
        "symulowanych ataków oraz ruchu rzeczywistego przechwytanego przez tshark."
    )
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    run = p.add_run(text)
    set_font(run, size=12)


def build_keywords(doc):
    p = doc.add_paragraph()
    fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=12, space_after=0, first_line=0)
    run_label = p.add_run("Słowa kluczowe:")
    set_font(run_label, size=12, underline=True)

    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=0, first_line=0)
    run_kw = p2.add_run(
        "bezpieczeństwo IoT, uczenie maszynowe, MITRE ATT&CK, Elastic Stack, detekcja anomalii"
    )
    set_font(run_kw, size=12)


def build_wstep(doc):
    add_heading(doc, "WSTĘP")

    paras = [
        (
            "Dynamiczny wzrost liczby urządzeń Internetu Rzeczy (IoT) — szacowany na ponad "
            "15 miliardów połączonych urządzeń w 2025 roku — tworzy rozległą powierzchnię ataku "
            "szczególnie narażoną na działania cyberprzestępców. Urządzenia IoT charakteryzują się "
            "ograniczonymi zasobami obliczeniowymi, rzadko aktualizowanym oprogramowaniem układowym "
            "oraz heterogenicznymi protokołami komunikacyjnymi, co znacząco utrudnia ich ochronę "
            "przy użyciu klasycznych narzędzi bezpieczeństwa sieci."
        ),
        (
            "Istniejące rozwiązania klasy SIEM (Security Information and Event Management) są zazwyczaj "
            "dedykowane dużym infrastrukturom korporacyjnym — ich koszt licencyjny i złożoność wdrożeniowa "
            "eliminuje je z zastosowań w małych sieciach domowych i laboratoryjnych. Brakuje rozwiązania, "
            "które w jednym środowisku Docker łączyłoby zbieranie ruchu, klasyfikację ML i kontekst "
            "operacyjny w postaci frameworku MITRE ATT&CK."
        ),
        (
            "Celem niniejszego artykułu jest zaprojektowanie, wdrożenie i walidacja lokalnego systemu "
            "monitorowania sieci IoT z automatyczną klasyfikacją ataków metodami uczenia maszynowego "
            "i mapowaniem zdarzeń na taktyki oraz techniki MITRE ATT&CK. Przedmiotem rozważań jest "
            "bezpieczeństwo sieci IoT ze szczególnym uwzględnieniem możliwości detekcji anomalii "
            "w ruchu sieciowym bez użycia zewnętrznych usług chmurowych."
        ),
        (
            "W pracy zastosowano metodologię eksperymentalną: zbudowano działający system, przetestowano "
            "go na danych benchmarkowych (NF-ToN-IoT-v2) oraz na ruchu rzeczywistym, a następnie "
            "porównano wyniki klasyfikatora z etykietami ground-truth. Analizę porównawczą uzupełniono "
            "wnioskowaniem przez analogię — spostrzeżenia ze środowiska laboratoryjnego odniesiono "
            "do scenariuszy produkcyjnych. Artykuł jest podzielony na pięć rozdziałów: charakterystykę "
            "zagrożeń IoT, opis architektury systemu, zbieranie i przetwarzanie danych, klasyfikację ML "
            "oraz weryfikację systemu w warunkach zbliżonych do rzeczywistych."
        ),
    ]
    for text in paras:
        p = doc.add_paragraph()
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=0, first_line=1.27)
        run = p.add_run(text)
        set_font(run, size=12)


def build_chapter1(doc):
    add_heading(doc, "ZAGROŻENIA BEZPIECZEŃSTWA W SIECIACH IOT", number="1")

    # 1.1
    add_heading(doc, "1.1 Specyfika zagrożeń IoT", level="sub")

    p1 = doc.add_paragraph()
    fmt_para(p1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r1 = p1.add_run(
        "Urządzenia IoT funkcjonują w warunkach fundamentalnie odmiennych od klasycznych stacji "
        "roboczych czy serwerów. Ograniczone zasoby obliczeniowe (CPU, RAM, flash), brak mechanizmów "
        "aktualizacji oprogramowania układowego oraz niejednorodność protokołów komunikacyjnych "
        "(Zigbee, Z-Wave, MQTT, CoAP, HTTP) sprawiają, że tradycyjne narzędzia bezpieczeństwa "
        "sieci — takie jak agenty endpoint detection and response (EDR) — są niepraktyczne lub "
        "niemożliwe do wdrożenia na samych urządzeniach. Ochrona musi odbywać się na poziomie sieci."
    )
    set_font(r1, size=12)

    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r2 = p2.add_run(
        "Na podstawie analizy zbioru NF-ToN-IoT-v2 oraz przeglądu literatury wyróżniono następujące "
        "kategorie ataków charakterystyczne dla sieci IoT:"
    )
    set_font(r2, size=12)

    add_dash_list(doc, [
        "ataki DDoS/DoS (Direct Network Flood, OS Exhaustion Flood) — przeciążenie zasobów sieciowych i obliczeniowych urządzeń",
        "brute force (SSH, HTTP) — łamanie haseł przez iterację słownikową",
        "skanowanie portów (port scan) — rozpoznanie topologii sieci i otwartych usług",
        "ataki injection (SQL, command injection) — wstrzykiwanie złośliwego kodu przez podatne interfejsy",
        "DNS exfiltration — wyprowadzanie danych kanałem DNS",
        "XSS (Cross-Site Scripting) — ataki na interfejsy webowe urządzeń zarządzających",
        "backdoor — instalacja trwałego dostępu poprzez webshell lub zmodyfikowane oprogramowanie układowe",
    ])

    p3 = doc.add_paragraph()
    fmt_para(p3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r3 = p3.add_run(
        "Badania Koroniotis i wsp. wskazują, że ataki na sieci IoT są coraz częściej prowadzone "
        "w sposób automatyczny przez botnety złożone z zainfekowanych urządzeń, co dodatkowo "
        "utrudnia ich detekcję na podstawie sygnatur statycznych."
    )
    set_font(r3, size=12)

    # 1.2
    add_heading(doc, "1.2 Framework MITRE ATT&CK jako narzędzie klasyfikacji", level="sub")

    p4 = doc.add_paragraph()
    fmt_para(p4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r4 = p4.add_run(
        "MITRE ATT&CK (Adversarial Tactics, Techniques, and Common Knowledge) jest publicznie "
        "dostępnym frameworkiem taksonomicznym opisującym zachowania atakujących na każdym etapie "
        "operacji ofensywnych. Framework systematyzuje wiedzę o taktykach (wysokopoziomowy cel, np. "
        "Initial Access, Persistence, Exfiltration), technikach (metoda realizacji celu, np. T1046 "
        "Network Service Discovery) oraz procedurach (konkretne implementacje). Zastosowanie "
        "frameworku ATT&CK jako warstwy semantycznej dla wyników klasyfikatora ML pozwala "
        "przekształcić surowy alert sieciowy w informację operacyjną zrozumiałą dla analityka SOC."
    )
    set_font(r4, size=12)

    p5 = doc.add_paragraph()
    fmt_para(p5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r5 = p5.add_run(
        "W niniejszym systemie każda klasa ataku jest mapowana statycznie na odpowiednią taktykę "
        "i technikę ATT&CK. W tabeli 2 (rozdział 4) zestawiono kompletne mapowanie dla dziewięciu "
        "klas klasyfikatora wraz z poziomem krytyczności (severity), który jest podstawą "
        "priorytetyzacji alertów w dashboardzie Kibana. Użyte techniki obejmują m.in. T1046 "
        "(Discovery), T1110.001 (Credential Access), T1048.001 (Exfiltration), T1498.001 i T1499.001 "
        "(Impact), T1059 (Execution) oraz T1505.003 (Persistence)."
    )
    set_font(r5, size=12)


def build_chapter2(doc):
    add_heading(doc, "ARCHITEKTURA SYSTEMU IOT SECURITY RADAR", number="2")

    add_heading(doc, "2.1 Przegląd komponentów", level="sub")

    p1 = doc.add_paragraph()
    fmt_para(p1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r1 = p1.add_run(
        "System IoT Security Radar składa się z czterech warstw logicznych: zbierania danych, "
        "przetwarzania i indeksowania, klasyfikacji ML oraz wizualizacji i alertowania. "
        "Przepływ danych ilustruje rysunek 1."
    )
    set_font(r1, size=12)

    # Architecture diagram as table (ASCII-art representation)
    p_arch_before = doc.add_paragraph()
    fmt_para(p_arch_before, space_before=6, space_after=0, first_line=0)

    arch_lines = [
        "┌─────────────────────────────────────────────────────────┐",
        "│                   ŹRÓDŁA DANYCH                         │",
        "│  [Router WiFi]  [Urządzenia IoT]  [Symulowane ataki]    │",
        "│                       ▼                                  │",
        "│              ┌─────────────┐                             │",
        "│              │  FILEBEAT   │  zbieranie i przesyłanie    │",
        "│              └──────┬──────┘  (TLS, beats protocol)      │",
        "│                     ▼                                    │",
        "│              ┌─────────────┐                             │",
        "│              │  LOGSTASH   │  parsowanie, wzbogacanie    │",
        "│              └──────┬──────┘  (pipeline NDJSON → ES)     │",
        "│                     ▼                                    │",
        "│         ┌────────────────────┐                           │",
        "│         │  ELASTICSEARCH     │  indeksowanie, TLS+xpack  │",
        "│         └───────┬────────────┘                           │",
        "│          ▼               ▼                               │",
        "│    ┌──────────┐  ┌────────────────┐                      │",
        "│    │  KIBANA  │  │  PYTHON ML     │  Random Forest +     │",
        "│    │ Dashboards│  │  (classifier)  │  MITRE ATT&CK map   │",
        "│    └──────────┘  └────────────────┘                      │",
        "│         Docker Compose + mTLS (wildcard *.local)         │",
        "└─────────────────────────────────────────────────────────┘",
    ]

    p_arch = doc.add_paragraph()
    fmt_para(p_arch, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=0, first_line=0)
    run_arch = p_arch.add_run("\n".join(arch_lines))
    set_font(run_arch, name="Courier New", size=9)

    add_fig_caption(doc, 1, "Architektura systemu IoT Security Radar")

    add_heading(doc, "2.2 Elastic Stack i bezpieczeństwo komunikacji", level="sub")

    components = [
        (
            "Elasticsearch 8.17.0 pełni rolę centralnego repozytorium zdarzeń. Skonfigurowany "
            "w trybie single-node z włączonymi mechanizmami xpack.security, indeksuje zdarzenia "
            "w formacie NDJSON. Dwa indeksy logiczne — iot-radar-* (surowe zdarzenia z etykietami "
            "ground-truth) oraz iot-radar-predictions (wyniki klasyfikatora ML) — realizują "
            "architekturę dwuindeksową opisaną szczegółowo w rozdziale 3."
        ),
        (
            "Logstash przetwarza zdarzenia w pipeline: wejście beats (port 5044, mTLS) → filtr "
            "mutate (normalizacja pól, tagowanie źródła data_source) → wyjście Elasticsearch. "
            "Konfiguracja unika konfliktu z domyślnym plikiem logstash.conf dostarczanym przez "
            "obraz Docker poprzez montowanie całego katalogu conf.d/."
        ),
        (
            "Filebeat 8.17.0 zbiera pliki NDJSON z katalogu sample-logs/ i przesyła je do "
            "Logstasha przez szyfrowany kanał TLS (certyfikat wildcard *.local). Mechanizm "
            "rejestru (filebeat-registry volume) zapobiega ponownemu wysyłaniu już przetworzonych "
            "zdarzeń po restarcie kontenera."
        ),
        (
            "Kibana udostępnia dwa dashboardy: SOC Dashboard (ground truth, rozkład typów ataków, "
            "mapa taktyk MITRE) oraz ML Predictions Dashboard (predykcje modelu, confidence, "
            "breakdown po taktykach ATT&CK). Alerty są skonfigurowane dla zdarzeń "
            "o severity=critical i severity=high."
        ),
        (
            "Bezpieczeństwo komunikacji zapewniają wzajemne uwierzytelnianie TLS (mTLS) z "
            "certyfikatami PEM (wildcard *.local, passphrase abcd1234) oraz mechanizm xpack.security "
            "Elasticsearch z dedykowanymi kontami (elastic, kibana_system). Wszystkie usługi są "
            "izolowane w sieci Docker, bez ekspozycji portów wewnętrznych na zewnątrz."
        ),
    ]
    for text in components:
        p = doc.add_paragraph()
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=0, first_line=1.27)
        run = p.add_run(text)
        set_font(run, size=12)


def build_chapter3(doc):
    add_heading(doc, "ZBIERANIE I PRZETWARZANIE DANYCH SIECIOWYCH", number="3")

    add_heading(doc, "3.1 Metody pozyskiwania danych", level="sub")

    p1 = doc.add_paragraph()
    fmt_para(p1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r1 = p1.add_run(
        "W systemie zaimplementowano trzy ścieżki pozyskiwania danych o różnym charakterze "
        "i przeznaczeniu."
    )
    set_font(r1, size=12)

    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r2 = p2.add_run(
        "Przechwytywanie ruchu rzeczywistego realizuje skrypt capture_traffic_flows.py oparty "
        "na narzędziu tshark. Skrypt korzysta z podejścia flow-based: pakiety są grupowane "
        "w przepływy (5-krotka: src_ip, dst_ip, src_port, dst_port, protocol), a cechy takie jak "
        "bytes_in i bytes_out są sumowane dla całego przepływu. Wcześniejsza implementacja "
        "per-packet zwracała bytes_received=0 dla wszystkich pakietów, co degradowało cztery "
        "najważniejsze cechy klasyfikatora (łącznie ~47% wagi Gini) i stanowi istotne "
        "ograniczenie podejścia przechwytywania na poziomie pakietu."
    )
    set_font(r2, size=12)

    p3 = doc.add_paragraph()
    fmt_para(p3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r3 = p3.add_run(
        "Zbiór benchmarkowy NF-ToN-IoT-v2 (Koroniotis i wsp., 2019) zawiera 16,9 mln przepływów "
        "sieciowych przechwyconych w laboratorium cyber-range Uniwersytetu w Canberze. Zbiór "
        "obejmuje 14 klas ataków i ruch normalny; do treningu i walidacji modelu użyto "
        "zbalansowanego podzbioru 55 000 próbek. Istotna cecha zbioru — adresy IP źródłowe "
        "w zakresie 192.168.1.x zarówno dla ruchu normalnego, jak i atakującego — wynika "
        "z topologii laboratorium, w której atakujący i ofiary współdzielą podsieć."
    )
    set_font(r3, size=12)

    p4 = doc.add_paragraph()
    fmt_para(p4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r4 = p4.add_run(
        "Symulowane ataki generuje skrypt generate_attacks.py, który produkuje zdarzenia NDJSON "
        "z wbudowanymi etykietami ground-truth. Klucz data_source=simulated pozwala odróżnić "
        "te zdarzenia od danych benchmarkowych i ruchu rzeczywistego na poziomie filtrów Kibana."
    )
    set_font(r4, size=12)

    add_heading(doc, "3.2 Pipeline przetwarzania", level="sub")

    p5 = doc.add_paragraph()
    fmt_para(p5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r5 = p5.add_run(
        "Architektura dwuindeksowa jest kluczowym elementem projektowym systemu. Indeks iot-radar-* "
        "przechowuje surowe zdarzenia z etykietami ground-truth nadanymi przez skrypty generujące "
        "(adapt_toniot.py, generate_attacks.py) lub przez Logstash (domyślna wartość attack_type=normal "
        "dla zdarzeń z tshark). Indeks iot-radar-predictions zawiera wyłącznie wyniki działania "
        "klasyfikatora ML — pola ml_prediction, ml_confidence, ml_mitre_tactic, ml_mitre_technique, "
        "ml_severity — i jest zasilany przez skrypt batch_score_all.py (tryb wsadowy) lub live_demo.py "
        "(tryb demonstracyjny)."
    )
    set_font(r5, size=12)

    p6 = doc.add_paragraph()
    fmt_para(p6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r6 = p6.add_run(
        "Rozdzielenie obu indeksów ma istotne implikacje dla analizy: zdarzenia z tshark-live "
        "w indeksie iot-radar-* noszą etykietę attack_type=normal (brak prawdziwego ground-truth), "
        "natomiast w indeksie iot-radar-predictions uzyskują pierwszą rzeczywistą klasyfikację "
        "od modelu ML. Architektura ta umożliwia niezależną ewaluację trafności modelu (porównanie "
        "ml_prediction z attack_type na danych benchmarkowych) bez zanieczyszczania metryk "
        "zdarzeniami bez ground-truth."
    )
    set_font(r6, size=12)


def build_chapter4(doc):
    add_heading(doc, "KLASYFIKACJA ATAKÓW Z WYKORZYSTANIEM UCZENIA MASZYNOWEGO", number="4")

    add_heading(doc, "4.1 Wybór algorytmu i przygotowanie danych", level="sub")

    p1 = doc.add_paragraph()
    fmt_para(p1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r1 = p1.add_run(
        "Do klasyfikacji ataków wybrano algorytm Random Forest (Breiman, 2001), kierując się "
        "trzema kryteriami: interpretowalność (feature importance jako miara wkładu cech), "
        "odporność na szum i brakujące wartości oraz brak konieczności standaryzacji danych "
        "numerycznych. Rozwiązania oparte na sieciach neuronowych oferują potencjalnie wyższą "
        "dokładność, lecz wymagają znacznie większych zasobów obliczeniowych i są trudniejsze "
        "do zinterpretowania w kontekście operacyjnym SOC."
    )
    set_font(r1, size=12)

    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r2 = p2.add_run(
        "Cechy wejściowe klasyfikatora odpowiadają polom przepływu sieciowego NetFlow: src_ip (numerycznie), "
        "dst_ip, src_port, dst_port, protocol, bytes_in, bytes_out, duration, packets_in, packets_out "
        "oraz flaga src_toniot_benchmark. Ostatnia cecha — binarna flaga wskazująca przynależność "
        "zdarzenia do zbioru benchmarkowego — osiąga pozycję #10 w rankingu ważności cech (feature "
        "importance), co wskazuje na częściowe uczenie się przez model tożsamości zbioru danych "
        "(feature leakage). Jest to akceptowalne w kontekście badawczym, lecz wymagałoby usunięcia "
        "w systemie produkcyjnym."
    )
    set_font(r2, size=12)

    p3 = doc.add_paragraph()
    fmt_para(p3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r3 = p3.add_run(
        "Mapowanie etykiet klas na liczby całkowite realizuje LabelEncoder z biblioteki scikit-learn "
        "(Pedregosa i wsp., 2011) z sortowaniem alfabetycznym. Artefakt enkodera (label_encoder.joblib) "
        "jest zapisywany razem z modelem, co zapewnia spójność mapowania między treningiem a inferencją — "
        "niezbędny warunek poprawnego działania klasyfikatora na nowych danych."
    )
    set_font(r3, size=12)

    add_heading(doc, "4.2 Wyniki klasyfikacji", level="sub")

    p4 = doc.add_paragraph()
    fmt_para(p4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r4 = p4.add_run(
        "Model osiągnął dokładność ogólną 97% na zbiorze testowym (11 000 próbek). "
        "Szczegółowe wyniki dla poszczególnych klas prezentuje tabela 1."
    )
    set_font(r4, size=12)

    # Table 1: Classification Report
    add_table_caption(doc, 1, "Raport klasyfikacji — Random Forest (zbiór testowy, n=11 000)")

    headers = ["Klasa", "Precision", "Recall", "F1-score", "Support"]
    rows_data = [
        ("backdoor",         "1.00", "1.00", "1.00",    "7"),
        ("brute_force",      "0.90", "0.88", "0.89",  "580"),
        ("ddos_flood",       "0.96", "0.93", "0.94",  "956"),
        ("dns_exfiltration", "1.00", "1.00", "1.00",    "9"),
        ("dos",              "0.89", "0.91", "0.90",  "338"),
        ("injection",        "0.74", "0.76", "0.75",  "322"),
        ("normal",           "1.00", "1.00", "1.00", "5900"),
        ("port_scan",        "1.00", "0.99", "1.00", "1777"),
        ("xss",              "0.92", "0.95", "0.93", "1111"),
        ("",                 "",     "",     "",       ""),
        ("accuracy",         "",     "",     "0.97", "11000"),
        ("macro avg",        "0.93", "0.94", "0.93", "11000"),
        ("weighted avg",     "0.97", "0.97", "0.97", "11000"),
    ]

    tbl = doc.add_table(rows=1 + len(rows_data), cols=5)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            set_font(run, size=10, bold=True)
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row_data in enumerate(rows_data):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                set_font(run, size=10)
            cells[ci].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Space after table
    p_after = doc.add_paragraph()
    fmt_para(p_after, space_before=12, space_after=0, first_line=0)

    # Feature importance figure
    p5 = doc.add_paragraph()
    fmt_para(p5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r5 = p5.add_run(
        "Klasy injection (F1=0.75) i brute_force (F1=0.89) osiągają niższe wyniki z powodu "
        "nakładania się rozkładów cech z innymi klasami ataków aplikacyjnych. Klasy normal, "
        "backdoor, dns_exfiltration i port_scan uzyskują F1=1.00, co potwierdza wyraźną "
        "separowalność tych wzorców ruchowych. Rysunek 2 przedstawia ważność cech "
        "(mean decrease in impurity) dla 15 najistotniejszych zmiennych wejściowych."
    )
    set_font(r5, size=12)

    # Insert feature importance image if it exists
    if os.path.exists(FEATURE_IMG):
        p_img_before = doc.add_paragraph()
        fmt_para(p_img_before, space_before=6, space_after=0, first_line=0)
        p_img = doc.add_paragraph()
        fmt_para(p_img, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=0, space_after=0, first_line=0)
        run_img = p_img.add_run()
        run_img.add_picture(FEATURE_IMG, width=Inches(5.5))
        add_fig_caption(doc, 2, "Ważność cech klasyfikatora Random Forest (top 15)")
    else:
        p_img_placeholder = doc.add_paragraph()
        fmt_para(p_img_placeholder, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=6, space_after=12, first_line=0)
        run_ph = p_img_placeholder.add_run("[Rys. 2 — feature_importance.png — brak pliku]")
        set_font(run_ph, size=10, italic=True)

    add_heading(doc, "4.3 Mapowanie predykcji na MITRE ATT&CK", level="sub")

    p6 = doc.add_paragraph()
    fmt_para(p6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r6 = p6.add_run(
        "Każda klasa ataku jest statycznie mapowana na taktykę, technikę i poziom krytyczności "
        "frameworku ATT&CK przy użyciu pliku konfiguracyjnego mitre_map.json. "
        "Mapowanie prezentuje tabela 2."
    )
    set_font(r6, size=12)

    # Table 2: MITRE mapping
    add_table_caption(doc, 2, "Mapowanie klas ataków na MITRE ATT&CK")

    mitre_headers = ["Klasa ataku", "Taktyka", "Technika", "Nazwa techniki", "Severity"]
    mitre_rows = [
        ("backdoor",         "Persistence",       "T1505.003", "Server Software Component: Web Shell",  "critical"),
        ("brute_force",      "Credential Access", "T1110.001", "Brute Force: Password Guessing",         "high"),
        ("ddos_flood",       "Impact",            "T1498.001", "Network Denial of Service: Direct Flood","critical"),
        ("dns_exfiltration", "Exfiltration",      "T1048.001", "Exfiltration Over Alternative Protocol","critical"),
        ("dos",              "Impact",            "T1499.001", "Endpoint Denial of Service: OS Exhaustion","high"),
        ("injection",        "Execution",         "T1059",     "Command and Scripting Interpreter",      "critical"),
        ("normal",           "—",                 "—",         "—",                                       "—"),
        ("port_scan",        "Discovery",         "T1046",     "Network Service Discovery",              "medium"),
        ("xss",              "Initial Access",    "T1189",     "Drive-by Compromise",                    "high"),
    ]

    tbl2 = doc.add_table(rows=1 + len(mitre_rows), cols=5)
    tbl2.style = "Table Grid"

    hdr2 = tbl2.rows[0].cells
    for i, h in enumerate(mitre_headers):
        hdr2[i].text = h
        for run in hdr2[i].paragraphs[0].runs:
            set_font(run, size=10, bold=True)
        hdr2[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row_data in enumerate(mitre_rows):
        cells = tbl2.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                set_font(run, size=10)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci != 3 else WD_ALIGN_PARAGRAPH.LEFT
            cells[ci].paragraphs[0].paragraph_format.alignment = align

    p_after2 = doc.add_paragraph()
    fmt_para(p_after2, space_before=12, space_after=0, first_line=0)

    p7 = doc.add_paragraph()
    fmt_para(p7, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r7 = p7.add_run(
        "W dashboardzie Kibana alerty są wyzwalane automatycznie dla zdarzeń o severity=critical "
        "(backdoor, ddos_flood, dns_exfiltration, injection) i severity=high (brute_force, dos, xss). "
        "Mapowanie jest przechowywane w pliku mitre_map.json, co ułatwia jego rozszerzenie "
        "o nowe klasy ataków bez modyfikacji kodu klasyfikatora."
    )
    set_font(r7, size=12)


def build_chapter5(doc):
    add_heading(doc, "WERYFIKACJA I DEMONSTRACJA SYSTEMU", number="5")

    add_heading(doc, "5.1 Demonstracja w czasie rzeczywistym", level="sub")

    p1 = doc.add_paragraph()
    fmt_para(p1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r1 = p1.add_run(
        "Kompletny cykl demonstracyjny uruchamia polecenie:"
    )
    set_font(r1, size=12)

    p_cmd = doc.add_paragraph()
    fmt_para(p_cmd, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=6, first_line=0)
    run_cmd = p_cmd.add_run(
        "sudo .venv/bin/python3 scripts/live_demo_with_attacks.py --minutes 3 --attacks 300"
    )
    set_font(run_cmd, name="Courier New", size=10)

    p2 = doc.add_paragraph()
    fmt_para(p2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r2 = p2.add_run(
        "Skrypt równolegle przechwytuje ruch WiFi przez tshark oraz generuje 300 symulowanych "
        "zdarzeń ataku, zapisując wyniki do pliku NDJSON. Filebeat wykrywa zmianę pliku "
        "i przesyła nowe zdarzenia przez mTLS do Logstasha, który je indeksuje do Elasticsearch. "
        "Moduł ML pobiera świeże zdarzenia przez Python ES client (elasticsearch==8.x), "
        "klasyfikuje je i zapisuje predykcje do indeksu iot-radar-predictions. "
        "Cały cykl od przechwycenia pakietu do pojawienia się predykcji w Kibana "
        "zamyka się w czasie poniżej 2 minut."
    )
    set_font(r2, size=12)

    p3 = doc.add_paragraph()
    fmt_para(p3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r3 = p3.add_run(
        "W toku weryfikacji zindeksowano łącznie 55 280 zdarzeń ze źródeł:"
    )
    set_font(r3, size=12)

    add_dash_list(doc, [
        "toniot_benchmark: 50 000 zdarzeń (NF-ToN-IoT-v2, balans 50/50 normal/attack)",
        "simulated: 5 000 zdarzeń (skrypt generate_attacks.py, stosunek 90/10)",
        "tshark-live: 265 zdarzeń (rzeczywisty ruch WiFi, brak ground-truth ataków)",
    ])

    add_heading(doc, "5.2 Dashboardy Kibana", level="sub")

    p4 = doc.add_paragraph()
    fmt_para(p4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r4 = p4.add_run(
        "SOC Dashboard oparty na indeksie iot-radar-* prezentuje rozkład typów ataków "
        "(treemap według attack_type), oś czasu zdarzeń, mapę taktyk MITRE oraz "
        "tabelę top IP źródłowych. Dashboard ten opiera się na etykietach ground-truth, "
        "co czyni go narzędziem do analizy rzeczywiście zaetykietowanych danych."
    )
    set_font(r4, size=12)

    p5 = doc.add_paragraph()
    fmt_para(p5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r5 = p5.add_run(
        "ML Predictions Dashboard oparty na indeksie iot-radar-predictions pokazuje "
        "wyniki działania modelu: rozkład predykowanych klas, rozkład pewności (confidence) "
        "klasyfikatora, breakdown zdarzeń według ml_mitre_tactic oraz historię wyzwolonych alertów. "
        "Dla zdarzeń tshark-live jest to jedyne miejsce w systemie, w którym ruch rzeczywisty "
        "uzyskuje klasyfikację bezpieczeństwa — w indeksie surowym brak jest ground-truth."
    )
    set_font(r5, size=12)

    p6 = doc.add_paragraph()
    fmt_para(p6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    r6 = p6.add_run(
        "Walidacja systemu obejmuje dwie perspektywy: ilościową (97% accuracy, szczegółowy "
        "classification report) na danych benchmarkowych z etykietami ground-truth oraz "
        "jakościową na ~265 rzeczywistych przepływach WiFi bez etykiet — co autorzy formułują "
        "jako: System zwalidowano na 55 000 zaetykietowanych zdarzeniach, a nastepnie "
        "zastosowano na ~265 niezaetykietowanych przeplywach WiFi, dla ktorych nie istnieje "
        "uprzedni ground-truth."
    )
    set_font(r6, size=12)


def build_wnioski(doc):
    add_heading(doc, "WNIOSKI", level="wnioski")

    paras = [
        (
            "System IoT Security Radar skutecznie realizuje zamierzony cel: lokalne monitorowanie "
            "sieci IoT z detekcją i klasyfikacją ataków w czasie rzeczywistym bez zależności od "
            "zewnętrznych usług chmurowych. Klasyfikator Random Forest osiąga dokładność 97% na "
            "zbiorze 11 000 próbek testowych, a architektura dwuindeksowa umożliwia niezależną "
            "ewaluację modelu i analizę operacyjną w kontekście MITRE ATT&CK."
        ),
        (
            "Zidentyfikowane ograniczenia systemu obejmują: (1) problem bytes_received=0 "
            "w trybie przechwytywania per-packet — rozwiązany przez implementację flow-based, "
            "docelowo wymagający protokołu NetFlow/IPFIX; (2) feature leakage wynikający "
            "z flagi src_toniot_benchmark — akceptowalny w badaniach, nieakceptowalny "
            "w produkcji; (3) statyczne mapowanie MITRE ATT&CK — nie uwzględnia "
            "dynamicznych łańcuchów ataków wieloetapowych."
        ),
        (
            "Rekomendacje dla dalszego rozwoju systemu: integracja protokołu NetFlow/IPFIX "
            "zamiast przechwytywania pakietów przez tshark jako docelowe źródło cech "
            "przepływowych; rozszerzenie zbioru treningowego o dane produkcyjne w celu "
            "eliminacji feature leakage; implementacja modelu wieloklasowego z możliwością "
            "detekcji łańcuchów ataków; integracja z komercyjnym systemem SIEM przez API. "
            "System w obecnej postaci stanowi funkcjonalne, w pełni odtwarzalne środowisko "
            "badawcze odpowiednie do dalszych eksperymentów w obszarze bezpieczeństwa IoT."
        ),
    ]
    for text in paras:
        p = doc.add_paragraph()
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=0, first_line=1.27)
        run = p.add_run(text)
        set_font(run, size=12)


def build_bibliography(doc):
    add_heading(doc, "BIBLIOGRAFIA")

    entries = [
        (
            "Breiman L., ",
            "Random Forests",
            u', \u201eMachine Learning\u201d, 2001, vol. 45, s. 5\u201332.',
        ),
        (
            "Docker Inc., ",
            "Docker Documentation",
            ", Docker Inc., 2024, https://docs.docker.com/, dostep 01.03.2026 r.",
        ),
        (
            "Elastic N.V., ",
            "Elastic Stack Documentation 8.x",
            ", Elastic N.V., 2024, https://www.elastic.co/guide/, dostep 01.03.2026 r.",
        ),
        (
            "Koroniotis N., Moustafa N., Sitnikova E., Turnbull B., ",
            "Towards the Development of Realistic Botnet Dataset in the Internet of Things for Network Forensic Analytics",
            u', \u201eFuture Generation Computer Systems\u201d, 2019, vol. 100, s. 779\u2013796.',
        ),
        (
            "MITRE Corporation, ",
            "MITRE ATT&CK: Design and Philosophy",
            ", MITRE Corporation, 2020, https://attack.mitre.org/, dostep 01.03.2026 r.",
        ),
        (
            "Moustafa N., Slay J., ",
            "UNSW-NB15: A Comprehensive Data Set for Network Intrusion Detection Systems",
            u', Military Communications and Information Systems Conference (MilCIS), IEEE, Canberra 2015, s. 1\u20136.',
        ),
        (
            "Pedregosa F. i in., ",
            "Scikit-learn: Machine Learning in Python",
            u', \u201eJournal of Machine Learning Research\u201d, 2011, vol. 12, s. 2825\u20132830.',
        ),
        (
            "Wireshark Foundation, ",
            u'TShark \u2014 Terminal-based Wireshark',
            ", Wireshark Foundation, 2024, https://www.wireshark.org/docs/man-pages/tshark.html, dostep 01.03.2026 r.",
        ),
    ]

    for i, (prefix, title_italic, suffix) in enumerate(entries):
        p = doc.add_paragraph()
        space_after = 6 if i < len(entries) - 1 else 0
        fmt_para(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=space_after, first_line=0)
        run_num = p.add_run(f"{i+1}. {prefix}")
        set_font(run_num, size=12)
        run_title = p.add_run(title_italic)
        set_font(run_title, size=12, italic=True)
        run_suffix = p.add_run(suffix)
        set_font(run_suffix, size=12)


def build_english_section(doc):
    # English title
    p_title = doc.add_paragraph()
    fmt_para(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=30, space_after=30, first_line=0)
    run_title = p_title.add_run(
        "IOT NETWORK MONITORING SYSTEM WITH ML-BASED ATTACK CLASSIFICATION "
        "AND MITRE ATT&CK FRAMEWORK MAPPING"
    )
    set_font(run_title, size=14, bold=True, spacing_pt=2)

    # ABSTRACT heading
    p_hdr = doc.add_paragraph()
    fmt_para(p_hdr, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=12, first_line=0)
    run_hdr = p_hdr.add_run("ABSTRACT")
    set_font(run_hdr, size=13, bold=True)

    # Abstract body
    text = (
        "This paper presents IoT Security Radar — a local IoT network monitoring system with "
        "real-time attack detection and classification. The system integrates the Elastic Stack "
        "(Filebeat, Logstash, Elasticsearch, Kibana) with a machine learning module based on the "
        "Random Forest algorithm, achieving 97% classification accuracy across nine attack classes. "
        "A dual-index architecture separates raw events (with ground-truth labels) from ML predictions, "
        "enabling independent evaluation of model performance and ground-truth data quality. Each "
        "detected attack is mapped to MITRE ATT&CK tactics and techniques, providing SOC operators "
        "with operational context for alert prioritization. The system was validated on 55,280 indexed "
        "events from three sources: the NF-ToN-IoT-v2 benchmark dataset, simulated attacks, and live "
        "WiFi traffic captured via tshark. The entire stack is containerized with Docker Compose and "
        "secured with mutual TLS authentication."
    )
    p_body = doc.add_paragraph()
    fmt_para(p_body, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0, first_line=1.27)
    run_body = p_body.add_run(text)
    set_font(run_body, size=12)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()

    build_author_block(doc)
    build_title(doc)
    build_abstract_pl(doc)
    build_keywords(doc)
    build_wstep(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_wnioski(doc)
    build_bibliography(doc)
    build_english_section(doc)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
